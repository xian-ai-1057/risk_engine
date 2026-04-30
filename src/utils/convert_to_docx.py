"""將 txt_v4 資料夾下的財報分析結果 TXT 轉換為 Word 文件。

每個 TXT 檔案包含 2 或 4 個 JSON 段落（以空行分隔）：
  - 2 段落：單一敘述、單一風險
  - 4 段落：合併敘述、合併風險、單一敘述、單一風險

輸出至 output/docx_v4/ 目錄，每個 TXT 對應一個 .docx。

Usage:
    python convert_to_docx.py
"""
import json
import re
import sys
from pathlib import Path

from docx import document as docx_document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
import docx


# ── 常數 ──────────────────────────────────────────

INPUT_DIR = Path(r"D:\02_專案\03_法金報告生成\01_Code\risk_analyzer\V6\output\v4\txt\txt_v4_新測試案例")  # 來源 TXT 目錄
OUTPUT_DIR = Path(r"D:\02_專案\03_法金報告生成\01_Code\risk_analyzer\V6\output\v4\docx\docx_v4_新測試案例")  # 輸出 DOCX 目錄

SECTION_NAMES: dict[str, str] = {
    "4-1": "財務結構",
    "4-2": "償債能力",
    "4-3": "經營效能",
    "4-4": "獲利能力",
    "4-5": "現金流量",
}

SECTION_KEYS = ["4-1", "4-2", "4-3", "4-4", "4-5"]

# 字型設定
FONT_NAME = "微軟正黑體"
FONT_NAME_EN = "Calibri"


# ── 解析 ──────────────────────────────────────────

def parse_txt_file(filepath: Path) -> list[dict[str, str]]:
    """解析 TXT 檔案，回傳 JSON 區塊列表。

    以空行分隔的多個 JSON 物件，逐一解析。

    Args:
        filepath: TXT 檔案路徑。

    Returns:
        包含 2 或 4 個 dict 的 list。
    """
    text = filepath.read_text(encoding="utf-8")

    # 用正規表達式找出所有 JSON 物件
    blocks = []
    # 匹配最外層的 { ... }
    depth = 0
    start = None
    for i, ch in enumerate(text):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start is not None:
                block_str = text[start:i + 1]
                blocks.append(json.loads(block_str))
                start = None

    return blocks


def extract_company_name(filename: str) -> str:
    """從檔名擷取公司名稱。

    檔名格式: {統編}_{公司名}_result.txt

    Args:
        filename: 檔案名稱。

    Returns:
        公司名稱字串。
    """
    match = re.match(r"\d+_(.+)_result\.txt", filename)
    if match:
        return match.group(1)
    return filename.replace("_result.txt", "")


# ── Word 生成 ─────────────────────────────────────

def _set_run_font(run: docx.text.run.Run, size_pt: int, bold: bool = False) -> None:
    """設定 run 的字型屬性。"""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # 設定東亞字型
    run.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        FONT_NAME,
    )


def _add_heading(doc: docx_document.Document, text: str, level: int) -> None:
    """新增標題段落並設定字型。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)

    size_map = {1: 18, 2: 14, 3: 12}
    _set_run_font(run, size_map.get(level, 12), bold=True)

    # 設定標題顏色為深色
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_body_paragraph(doc: docx_document.Document, text: str) -> None:
    """新增正文段落並設定格式。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, 11)

    # 段落格式：1.5 倍行距、段後 6pt
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(6)


def _write_paired_sections(
    doc: docx_document.Document,
    group_title: str,
    narrative_data: dict[str, str],
    risk_data: dict[str, str],
) -> None:
    """寫入一組配對區塊（每個 section 的敘述與風險放在一起）。

    Args:
        doc: Word 文件物件。
        group_title: 群組標題（如「合併」、「單一」、「財報分析」）。
        narrative_data: 財報敘述資料，包含 "4-1"~"4-5" key。
        risk_data: 風險評估資料，包含 "4-1"~"4-5" key。
    """
    _add_heading(doc, group_title, level=2)

    for key in SECTION_KEYS:
        narrative = narrative_data.get(key, "")
        risk = risk_data.get(key, "")
        if not narrative and not risk:
            continue
        section_label = f"{key} {SECTION_NAMES[key]}"
        _add_heading(doc, section_label, level=3)
        if narrative:
            _add_body_paragraph(doc, narrative)
        if risk:
            _add_body_paragraph(doc, risk)


def convert_single_file(txt_path: Path, output_path: Path) -> bool:
    """將單一 TXT 檔案轉換為 Word 文件。

    Args:
        txt_path: 來源 TXT 路徑。
        output_path: 輸出 DOCX 路徑。

    Returns:
        轉換是否成功。
    """
    blocks = parse_txt_file(txt_path)
    n = len(blocks)
    if n not in (2, 4):
        if n == 1:
            print(f"  ⚠ 跳過（僅 1 個區塊，缺少風險評估段落）: {txt_path.name}")
        elif n == 3:
            print(
                f"  ⚠ 跳過（僅 3 個區塊，預期 4 個：合併敘述/合併風險/單一敘述/單一風險，"
                f"缺少單一風險段落）: {txt_path.name}"
            )
        else:
            print(f"  ⚠ 跳過（區塊數 {n}，預期 2 或 4）: {txt_path.name}")
        return False

    company = extract_company_name(txt_path.name)
    doc = docx.Document()

    # 文件標題
    _add_heading(doc, company, level=1)

    if len(blocks) == 4:
        # 合併敘述、合併風險、單一敘述、單一風險
        _write_paired_sections(doc, "合併", blocks[0], blocks[1])
        _write_paired_sections(doc, "單一", blocks[2], blocks[3])
    else:
        # 單一敘述、單一風險
        _write_paired_sections(doc, "財報分析", blocks[0], blocks[1])

    doc.save(str(output_path))
    return True


# ── 主流程 ─────────────────────────────────────────

def main() -> None:
    """批次轉換所有 TXT 檔案為 Word 文件。"""
    if not INPUT_DIR.is_dir():
        print(f"錯誤: 找不到輸入目錄 {INPUT_DIR}")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    txt_files = sorted(INPUT_DIR.glob("*_result.txt"))
    if not txt_files:
        print(f"找不到任何 _result.txt 檔案於 {INPUT_DIR}")
        sys.exit(1)

    print(f"找到 {len(txt_files)} 個檔案，開始轉換...\n")

    success = 0
    failed = 0
    for txt_path in txt_files:
        docx_name = txt_path.stem + ".docx"
        output_path = OUTPUT_DIR / docx_name
        company = extract_company_name(txt_path.name)

        if convert_single_file(txt_path, output_path):
            print(f"  ✓ {company}")
            success += 1
        else:
            failed += 1

    print(f"\n完成: {success} 個成功, {failed} 個失敗")
    print(f"輸出目錄: {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
